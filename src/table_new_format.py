"""
Author : Daniel Dekhtyar
Latest update : 2/10/2023

The `add_rows()` function is the main function of the script. It iterates through all tables in the document and performs the following steps for each table:
1. Adds 3 rows to the table and moves the text down by two rows.
2. Formats the table by merging cells, setting text values, and applying shading.
3. Sets the text and formatting of the first row of the table.
4. Styles the table by setting column widths, aligning paragraphs, and changing the font.
5. Saves the modified document with the changes.

The code modifies the tables in the Word document by adding 3 rows, moving the text down, and applying formatting. The modified document is then saved with the changes.

Example Usage:
```python
# Create a Word document object
document = Document()

# Create a table object
table = document.add_table(rows=5, cols=6)

# Call the add_rows() function to add 3 rows to the table and format it
add_rows()

# The table will have 8 rows in total, with the first 3 rows added and the remaining 5 rows from the original table
# The text in each cell will be moved down by two rows
# The table will be formatted with merged cells, shading, and specific text values in the first row
# The modified document will be saved with the changes
```
"""

"""
Text prints in revers bug solution : 
Before that only the first table was right and the rest were reversed.
I solved it by defining a global variable that keeps tack of the number of tables that is already modified.
If table_count = 0 means that we are still on the first table.
If it is bigger then 0 then we have passed the first table.
So to fix the bug it intentionally prints the values in reverse and the final product is as expected !
a.k.a backward logic
"""


from docx.api import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Pt
import time
start_time = time.time()

table_count = 0

# add_rows() is the main function of the script


def add_rows():
    """
    The function `add_rows()` iterates through all tables in a Word document, adds 3 rows to each table,
    moves text down the tables to the new format, applies required formatting to the tables, and saves
    the modified document.
    """
    rahshal = Document(r"C:\Users\Daniel\Desktop\Iron Dome\רכשי לב.docx")
    global table_count
    # Iterate through all tables in the document
    # Add 3 rows and move text down tables to the new format
    for table in rahshal.tables:
        add_3_rows_and_move_text_down(table)
        make_table_format_as_required(table, rahshal)
        first_row(table)
        style_the_docx_table(table)
        table_count += 1
        print(
            f"Table {table_count} of {
                len(rahshal.tables)} modified according to the new format"
        )

    rahshal.save(r"C:\Users\Daniel\Desktop\Iron Dome\רכשי לב.docx")
    print("All done and save successfully !")
    print(f"--- The code took {time.time() - start_time} seconds to run ---")


def add_3_rows_and_move_text_down(table):
    """
    The function adds three new rows to a table and moves the text in each cell down by two rows.

    Args:
    table: The parameter 'table' is expected to be a data structure representing a table. It could be
    a list of lists, where each inner list represents a row and contains the values of the cells in that
    row. Alternatively, it could be an object representing a table with methods to add rows and access
    cells
    """
    # Implementation of the function
    pass
    # Create a list to store the original values
    original_values = []

    # Iterate through the existing rows and cells in the table starting from row 2
    for i in range(1, len(table.rows)):
        row_values = [cell.text for cell in table.rows[i].cells]
        original_values.append(row_values)

    # Insert three new rows
    for _ in range(4):
        table.add_row().cells

    # Iterate through the 'original_values' list and add the content two rows below
    for i, row_values in enumerate(original_values):
        new_row_index = i + 5  # Calculate the new row index
        if new_row_index < len(table.rows):
            for j, cell_value in enumerate(row_values):
                table.cell(new_row_index, j).text = cell_value


def make_table_format_as_required(table, rahshal):
    """
    The function `make_table_format_as_required` modifies a table by merging cells, setting text values,
    and applying shading.

    Args:
    table: The "table" parameter is a table object that represents a table in a document. It is used
    to access and modify the cells and properties of the table.
    rahshal: The parameter "rahshal" is not defined in the code snippet you provided. Please provide
    the definition of "rahshal" so that I can assist you further.
    """
    for i in range(4):
        for cell in table.rows[i].cells:
            cell.text = ""

    table.rows[4].cells[0].merge(table.rows[4].cells[1])  # Merge cells 1 and 2
    table.rows[4].cells[2].merge(table.rows[4].cells[3])  # Merge cells 3 and 4
    table.rows[4].cells[4].merge(table.rows[4].cells[5])  # Merge cells 5 and 6

    if table_count == 0:
        table.rows[4].cells[0].text = "WGS84 GEO DM"
        table.rows[4].cells[2].text = "WGS84 GEO D"
        table.rows[4].cells[4].text = "ED50 UTM36"
    else:
        table.rows[4].cells[4].text = "WGS84 GEO DM"
        table.rows[4].cells[2].text = "WGS84 GEO D"
        table.rows[4].cells[0].text = "ED50 UTM36"

    set_shading(0, 3, "ffffff", table)
    set_shading(4, 6, "bfbfbf", table)

    first_row = table.rows[0]
    table_element = table._tbl
    row_element = first_row._tr
    table_element.remove(row_element)
    merge_third_row(table, rahshal)


def set_shading(row_start: int, row_end: int, color: str, table):
    """
    The function `set_shading` sets the shading color for a range of cells in a table.

    Args:
    row_start (int): The starting row index where the shading should be applied.
    row_end (int): The `row_end` parameter is an integer that represents the ending row index
    (exclusive) for setting the shading. This means that the shading will be applied to all rows from
    `row_start` to `row_end - 1`.
    color (str): The "color" parameter is a string that represents the color of the shading that you
    want to apply to the cells in the table. It should be a valid color value, such as "FF0000" for red
    or "00FF00" for green.
    table: The `table` parameter is the table object that you want to apply shading to. It should be
    an instance of a table class, such as `docx.table.Table`.
    """
    for row in range(row_start, row_end):  # Set the shading
        for cell in range(len(table.rows[row].cells)):
            # GET CELLS XML ELEMENT
            cell_xml_element = table.rows[row].cells[cell]._tc
            # RETRIEVE THE TABLE CELL PROPERTIES
            table_cell_properties = cell_xml_element.get_or_add_tcPr()
            # CREATE SHADING OBJECT
            shade_obj = OxmlElement("w:shd")
            # SET THE SHADING OBJECT
            shade_obj.set(qn("w:fill"), color)
            # APPEND THE PROPERTIES TO THE TABLE CELL PROPERTIES
            table_cell_properties.append(shade_obj)


def merge_third_row(table, rahshal):
    """
    The function `merge_third_row` merges adjacent cells in the third row of a table and sets the text
    in the merged cell to "קואורדינטות" with a specific font size and style.

    Args:
    table: The "table" parameter is the table object that you want to modify. It should be an instance
    of the Table class.
    rahshal: The parameter "rahshal" is not used in the code provided. It seems to be an unused
    variable.
    """
    for cell in range(len(table.rows[2].cells) - 1):
        cell_a = table.cell(2, cell)
        cell_b = table.cell(2, cell + 1)
        merged_cell = cell_a.merge(cell_b)
        merged_cell.text = "קואורדינטות"
        for paragraph in merged_cell.paragraphs:
            for run in paragraph.runs:
                """You can enable different designs"""
                # run.font.bold = True
                # run.underline = True
                run.font.size = Pt(12)
                run.font.italic = True


def style_the_docx_table(table):
    """
    Styles a table in a Word document by setting the column width, aligning the paragraphs, and changing the font.

    Args:
        table (Table): The table object that needs to be styled.

    Returns:
        None. Modifies the table object directly.

    Example Usage:
        table = document.add_table(rows=3, cols=4)
        style_the_docx_table(table)

    Code Analysis:
        - The function sets the `autofit` property of the table to `False` to prevent the columns from automatically adjusting their width.
        - It iterates through each column in the table and sets the width of each column to 3 centimeters using the `width` property.
        - It then iterates through each cell in the table and applies the desired styling to the paragraphs and runs within each cell.
        - For each paragraph in a cell, it sets the alignment to center using the `alignment` property.
        - For each run in a paragraph, it changes the font to "Calibri Light" using the `font.name` property.
    """
    table.autofit = False
    # Define the desired column width in centimeters (3 cm)

    # Iterate through the columns and set their widths
    for column in range(len(table.columns)):
        col = table.columns[column]
        col.width = Cm(3.0)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.font.name = "Calibri Light"


def first_row(table):
    """
    Sets the text and formatting of the first row of a table and applies shading to the first row.

    Args:
        table (docx.table.Table): The table object that represents the table in which the first row needs to be modified.

    Returns:
        None

    Example Usage:
        # Create a table object
        table = docx.table.Table()

        # Call the first_row function to set the text and formatting of the first row
        first_row(table)

        # The first row of the table will have the specified text and formatting, and will be shaded with light gray
    """
    if table_count == 0:
        table.cell(0, 5).text = "שם אזור"
        table.cell(0, 4).text = "מדיניות הגנה"
        table.cell(0, 3).text = "עדיפות"
        table.cell(0, 2).text = "מדיניות הפעלה"
        table.cell(0, 1).text = "עדכון אחרון"
        table.cell(0, 0).text = "הערות"
    else:
        table.cell(0, 0).text = "שם אזור"
        table.cell(0, 1).text = "מדיניות הגנה"
        table.cell(0, 2).text = "עדיפות"
        table.cell(0, 3).text = "מדיניות הפעלה"
        table.cell(0, 4).text = "עדכון אחרון"
        table.cell(0, 5).text = "הערות"

    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.italic = True

    # Set the shading of row 0 to light gray
    set_shading(0, 1, "bfbfbf", table)


if __name__ == "__main__":
    add_rows()
