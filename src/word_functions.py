"""
This file contains functions for working with Word documents.
Those functions are called from xl_to_rahshal.py
"""


from docx import Document
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from termcolor import colored
import excel_functions


def copy_coordinates_from_xl_to_rahshal(
    nz_xl, docx_table: Document, xl_row: int, table_count: int
) -> None:
    """
    The function `copy_coordinates_from_xl_to_rahshal` copies coordinates from an Excel file to a
    specified table in a Word document.

    :param nz_xl: The variable `nz_xl` is likely referring to an Excel workbook or worksheet object that
    contains the coordinates data. It is used to access specific cells in the Excel sheet
    :param docx_table: The parameter `docx_table` is of type `Document`, which is likely referring to a
    document object in a word processing software such as Microsoft Word. This object represents a
    document that contains tables
    :type docx_table: Document
    :param xl_row: The `xl_row` parameter represents the row number in the Excel sheet from which the
    coordinates will be copied
    :type xl_row: int
    """
    # Start at row 5 because we need to leave space for the 5 default rows
    for row in range(5, len(docx_table.rows)):
        # Explanation for the if statement cn be found in xl_to_rahshal.py docstring
        if table_count == 0:
            for column in range(1, 7):  # Columns 1 to 7 in the excel
                cell = nz_xl.cell((xl_row + row) - 1, column + 1)
                # 'row + 2' because don't need to copy the first 2 rows
                if cell.value is not None:
                    docx_table.cell(row, column - 1).text = str(cell.value)
                    # 'column - 1' because in docx it start from index 0 and in excel it starts from index 1
        else:
            for column in range(1, 7):  # Columns 1 to 7 in the excel
                cell = nz_xl.cell((xl_row + row) - 1, column + 1)
                # 'row + 2' because don't need to copy the first 2 rows
                if cell.value is not None:
                    docx_table.cell(row, abs(column - 6)).text = str(cell.value)
                    # 'abs(column - 6)' explanation in the doc in the start of the document


def find_area_in_rahshal(
    area_name: str, rahshal: Document, table_count: int
) -> [int or None]:
    """
    The function `find_area_in_rahshal` takes an area name and a document object as input, and returns
    the index of the table containing the area name in the document, or None if the area is not found.

    :param area_name: The area name is a string that represents the name of the area you want to find in
    the "rahshal" document
    :type area_name: str
    :param rahshal: The parameter "rahshal" is of type "Document", which suggests that it is a document
    object, possibly representing a Microsoft Word document
    :type rahshal: Document
    :return: The function `find_area_in_rahshal` returns an integer value representing the index of the
    table in the `rahshal` document that contains the specified `area_name`. If the `area_name` is not
    found in the document, the function returns `None`.
    """
    for table_index, table in enumerate(rahshal.tables):
        # Explanation for the if statement cn be found in xl_to_rahshal.py docstring
        if table_count == 0:
            if table.cell(1, 5).text == area_name:
                return int(table_index)
        elif table.cell(1, 0).text == area_name:
            return int(table_index)
        table_index += 1
    # If area is not in the rahshal docx, this error message will appear
    print(colored("---------------------------------", "red", attrs=["bold"]))
    print(
        colored(f"The area {area_name[::-1]} is not in the docx", "red", attrs=["bold"])
    )
    print(colored("---------------------------------", "red", attrs=["bold"]))
    return None


def update_table_dimensions_in_rahshal(
    rahshal: Document, nz_xl, xl_row: int, table_index: int
) -> None:
    """
    The function `update_table_dimensions_in_rahshal` updates the dimensions of a table in a document by
    adding or removing rows based on the specified number of rows from an Excel file.

    :param rahshal: The `rahshal` parameter is a `Document` object, which represents a Word document
    :type rahshal: Document
    :param nz_xl: The parameter `nz_xl` is not defined in the given code snippet. It seems to be a
    missing variable or function that is required to determine the number of rows in the Excel table.
    Please provide more information or update the code snippet with the definition of `nz_xl` so that I
    :param xl_row: The `xl_row` parameter represents the row number in the Excel file
    :type xl_row: int
    :param table_index: The parameter `table_index` is the index of the table in the `rahshal` document
    that you want to update the dimensions for. It is used to access the specific table in the `tables`
    list
    :type table_index: int
    """
    tables = rahshal.tables
    docx_table = tables[table_index]
    rows_old_table: int = len(docx_table.rows)
    rows_new_table: int = excel_functions.xl_table_dimensions(nz_xl, xl_row) + 3
    # Add +3 to rows_new_table because 3 rows added to the docx_table to have the new format

    if rows_old_table == rows_new_table:
        pass
    elif rows_old_table < rows_new_table:
        rows_to_add = rows_new_table - rows_old_table
        for _ in range(rows_to_add):
            docx_table.add_row()
    elif rows_old_table > rows_new_table:
        rows_to_remove = rows_old_table - rows_new_table
        for _ in range(rows_to_remove):
            row = docx_table.rows[len(docx_table.rows) - 1]
            table_element = docx_table._tbl
            row_element = row._tr
            table_element.remove(row_element)


def style_the_docx_table(docx_table: Document) -> None:
    """
    The function `style_the_docx_table` applies specific formatting to a table in a Word document.

    :param docx_table: The parameter `docx_table` is of type `Document`, which represents a Word
    document in the python-docx library. It is assumed that `docx_table` contains a table that needs to
    be styled
    :type docx_table: Document
    """
    for row in docx_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.font.name = "Calibri Light"
