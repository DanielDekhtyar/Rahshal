from docx.api import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time


start_time = time.time()


def main():
    rahshal = Document(r"C:\Users\Daniel\Desktop\Iron Dome\רכשי לב.docx")
    table_count = 0
    # Iterate through all tables in the document
    for table in rahshal.tables:
      add_2_rows_and_move_text_down(table)
      make_table_look_as_required(table)
      style_the_docx_table(table)
      table_count += 1
      print(f"Table {table_count} of {len(rahshal.tables)}")

    rahshal.save(r"C:\Users\Daniel\Desktop\Iron Dome\רכשי לב.docx")
    print("All done and save successfully !")
    print(f"--- The code took {time.time() - start_time} seconds to run ---")


def add_2_rows_and_move_text_down(table):
    # Create a list to store the original values
    original_values = []

    # Iterate through the existing rows and cells in the table starting from row 2
    for i in range(1, len(table.rows)):
        row_values = []
        for cell in table.rows[i].cells:
            row_values.append(cell.text)
        original_values.append(row_values)

    # Insert two new rows at the top of the table
    for _ in range(2):
        new_row = table.add_row().cells

    # Iterate through the 'original_values' list and add the content two rows below
    for i, row_values in enumerate(original_values):
        new_row_index = i + 3  # Calculate the new row index
        if new_row_index < len(table.rows):
            for j, cell_value in enumerate(row_values):
                table.cell(new_row_index, j).text = cell_value


def merge_second_row(new_row2, cell):
    # Merge the cells in the second row into one cell
    new_cell2 = new_row2[0]
    for cell in new_row2[1:]:
        new_cell2.merge(cell)
    new_cell2.text = "קורדינטות"


def style_the_docx_table(docx_table):
    for row in docx_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


if __name__ == "__main__":
    main()
