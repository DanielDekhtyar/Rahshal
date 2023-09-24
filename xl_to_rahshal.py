import openpyxl
from docx import Document
import time

# The code copies the coordinates of a specific area from an excel file to a table in Microsoft Word file called 'רכשי לב', or for short 'rahshal'

start_time = time.time()


def main():
    # Load the excel workbook
    wb = openpyxl.load_workbook(r"C:\Users\Daniel\Desktop\Iron Dome\Coordinates.xlsx")
    nz_xl = wb.active  # nz means נ.צ
    rahshal = Document(r"C:\Users\Daniel\Desktop\Iron Dome\רכשי לב.docx")
    xl_row = 1
    table_index = 0
    while xl_row < nz_xl.max_row:
        area_name, xl_row = find_area_in_xl(nz_xl, xl_row)
        if area_name != " ":
            results = is_area_in_rahshal(area_name, rahshal)
            if results is not None:
                # Check if None is returned, meaning that the area is not in the docx
                area_name, paragraph = results
                docx_table = update_table_dimensions_in_rahshal(
                    rahshal, paragraph, nz_xl, xl_row, table_index
                )
                table_index += 1

    rahshal.save(r"C:\Users\Daniel\Desktop\Iron Dome\רכשי לב.docx")
    wb.close()
    print("All done and save successfully !")
    print(f"--- The code took {time.time() - start_time} seconds to run ---")


# 'TypeError: cannot unpack non-iterable NoneType object' solved by putting all the return values in to one tuple
# If the function didn't do it's work then return this one variable as None
# At the receiving end, check if the return value is None, else you can safely unpack the tuple and use it.


def find_area_in_xl(nz_xl, xl_row):  # TESTED AND DONE !
    for row in range(xl_row + 1, nz_xl.max_row + 1):
        cell_value = nz_xl[f"A{row}"].value
        if cell_value is not None:
            if "קורדינטות" not in cell_value:
                return cell_value, row
    return " ", row


def is_area_in_rahshal(area_name, rahshal):
    for i, paragraph in enumerate(rahshal.paragraphs):
        if area_name == paragraph.text:
            results = area_name, paragraph
            return results

    print(f"The area {area_name[::-1]} is not in the docx")
    # If area is not in rahshal, an error message will appear
    return None


def update_table_dimensions_in_rahshal(rahshal, paragraph, nz_xl, xl_row, table_index):
    tables = rahshal.tables
    docx_table = tables[table_index]
    rows_old_table = len(docx_table.rows)
    rows_new_table = xl_table_dimensions(nz_xl, xl_row)
    if rows_old_table < rows_new_table:
        rows_to_add = rows_new_table - rows_old_table
        for _ in range(rows_to_add):
            docx_table.add_row()
    elif rows_old_table > rows_new_table:
        rows_to_remove = rows_old_table - rows_new_table
        for _ in range(rows_to_remove):
            row = docx_table.rows[len(docx_table.rows) - 1]
            table_element = docx_table._tbl
            row_element = row._tr
            table_element.remove(row_element)
    elif rows_old_table == rows_new_table:
        pass

    return docx_table


def xl_table_dimensions(nz_xl, xl_row):  # TESTED AND DONE !
    int(xl_row)
    # Finds the number of rows in the table
    first_row_of_table = None
    last_row_of_table = None
    # I added +1 to max_row because otherwise it will go upto the one to last but not the last row
    max_row = nz_xl.max_row + 1

    for row in range(int(xl_row), max_row):
        cell_in_column_B = nz_xl.cell(row, 2)  # 2 corresponds to column B
        if cell_in_column_B.value is not None:
            first_row_of_table = row
            break

    for row in range(first_row_of_table, max_row):
        cell_in_column_B = nz_xl.cell(row, 2)  # 2 corresponds to column B
        if cell_in_column_B.value is None:
            last_row_of_table = row
            break
        elif row == max_row:
            last_row_of_table = max_row

    number_of_rows_in_table = last_row_of_table - first_row_of_table
    return number_of_rows_in_table


if __name__ == "__main__":
    main()
